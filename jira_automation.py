"""
JIRA to DOCX Automation Script

This script fetches JIRA issues from multiple projects, generates AI summaries,
extracts deliverables using AI analysis, and creates a comprehensive Word document report.

Author: AI Assistant
Date: July 2025
"""

import os
import requests
import json
from docx import Document
from dotenv import load_dotenv
from typing import List, Dict, Any
from datetime import datetime

# Load environment variables
load_dotenv()

# Configuration constants
JIRA_TOKEN = os.getenv("JIRA_TOKEN")
JIRA_URL = os.getenv("JIRA_URL")
JQL = os.getenv("JIRA_JQL")
OLLAMA_URL = "http://localhost:11434/api/generate"


class JiraToDocxAutomation:
    """
    Main automation class for JIRA to DOCX reporting
    
    This class handles:
    - Fetching issues from JIRA
    - Generating AI summaries
    - AI-powered deliverable extraction  
    - Word document generation
    """
    
    def __init__(self, project_names: List[str]):
        """
        Initialize the automation with project names
        
        Args:
            project_names: List of JIRA project names to process
        """
        self.validate_config()
        self.project_names = project_names if isinstance(project_names, list) else [project_names]
    
    @classmethod
    def from_env_projects(cls):
        """
        Create automation instance using project names from environment variable
        
        Expected format: JIRA_PROJECTS="Project1,Project2,Project3"
        
        Returns:
            List of project names from environment variable
            
        Raises:
            ValueError: If JIRA_PROJECTS is not set
        """
        project_names_env = os.getenv("JIRA_PROJECTS", "")
        if not project_names_env:
            raise ValueError("JIRA_PROJECTS must be set in .env file (comma-separated list)")

        project_names = [name.strip() for name in project_names_env.split(",")]
        return project_names
    
    def validate_config(self):
        """
        Validate that all required environment variables are set
        
        Raises:
            ValueError: If any required environment variable is missing or invalid
        """
        if not JIRA_TOKEN or JIRA_TOKEN == "your_token_here":
            raise ValueError("JIRA_TOKEN must be set in .env file")
        if not JIRA_URL or "yourdomain" in JIRA_URL:
            raise ValueError("JIRA_URL must be set to your actual JIRA instance in .env file")
        if not JQL:
            raise ValueError("JIRA_JQL must be set in .env file")

    def fetch_issues(self, project_name: str) -> List[Dict[str, Any]]:
        """
        Fetch issues from JIRA using JQL query
        
        Args:
            project_name: Name of the JIRA project
            
        Returns:
            List of issue dictionaries from JIRA API
        """
        try:
            headers = {
                "Authorization": f"Bearer {JIRA_TOKEN}",
                "Accept": "application/json",
            }
            
            params = {
                "jql": f"project = '{project_name}' {JQL}",
                "fields": "issuetype,key,summary,status,project,priority,assignee,reporter,created,updated,duedate",
            }

            print(f"Fetching issues from JIRA with JQL: project = '{project_name}' {JQL}")
            response = requests.get(
                "https://tracker.nci.nih.gov/rest/api/2/search", 
                headers=headers, 
                params=params,
                timeout=30
            )
            
            if response.status_code != 200:
                raise Exception(f"JIRA API error: {response.status_code} - {response.text}")
            
            data = response.json()
            issues = data.get("issues", [])
            print(f"Successfully fetched {len(issues)} issues from JIRA")
            return issues
            
        except requests.exceptions.RequestException as e:
            print(f"Error fetching issues from JIRA: {e}")
            return []
        except Exception as e:
            print(f"Unexpected error: {e}")
            return []
    
    def summarize_with_ollama(self, text: str) -> str:
        """
        Generate AI summary using Ollama LLM
        
        Args:
            text: The text content to summarize
            
        Returns:
            AI-generated summary text
        """
        try:
            # Prepare the request body for Ollama API
            body = {
                "model": "llama3",
                "prompt": (
                    "You are a project manager assistant. Given a list of JIRA issues or tasks with the fields: "
                    "Issue Type, Issue Key, Summary, and Status, create a concise and professional high-level summary "
                    "of completed, planned for this specific project in the current or upcoming month. "
                    "The overall summary should be limited to 150 words. "
                    "Do not include any specific issue keys in your summary. "
                    "Do not list individual issues. No explanation needed—just the summary.\n\n"
                    f"Here is the list of issues for this project: {text}"
                ),
                "stream": False
            }

            headers = {"Content-Type": "application/json"}
            
            print("Generating summary with Ollama...")
            response = requests.post(OLLAMA_URL, json=body, headers=headers)
            
            if response.status_code != 200:
                print(f"Ollama API error: {response.status_code} - {response.text}")
                return f"Error generating summary: {response.text}"
            
            result = response.json()
            summary = result.get("response", "")
            return summary.strip()
            
        except requests.exceptions.RequestException as e:
            print(f"Error connecting to Ollama: {e}")
            return f"Error connecting to Ollama: {e}"
        except Exception as e:
            print(f"Unexpected error during summarization: {e}")
            return f"Error generating summary: {e}"


    def generate_word_document(self, projects_data: Dict[str, Dict], filename: str = "JIRA_Summary_Report.docx"):
        """
        Generate comprehensive Word document with deliverables and project details
        
        Args:
            projects_data: Dictionary with project names as keys and project data as values
            filename: Output filename for the Word document
        """
        try:
            print(f"Generating Word document: {filename}")
            
            doc = Document()
            
            # Add document title
            title = doc.add_heading("Projects Monthly Status Report", 1)
            title.alignment = 1  # Center alignment
            
            # Generate deliverable overview table at the beginning
            project_deliverables = self.extract_deliverables(projects_data)
            if project_deliverables:
                self.generate_deliverable_table(doc, project_deliverables)
                doc.add_page_break()
            
            # Process each project section
            for project_name, project_data in projects_data.items():
                issues = project_data.get("issues", [])
                project_summary = project_data.get("ai_summary", "No summary available")
                
                # Add project heading
                doc.add_heading(f"{project_name}: Tasks completed or to be continued in the upcoming month.", 2)
                
                # Create issues table
                issues_table = doc.add_table(rows=1, cols=4)
                issues_table.style = 'Table Grid'
                
                # Set table headers
                hdr_cells = issues_table.rows[0].cells
                hdr_cells[0].text = 'Issue Type'
                hdr_cells[1].text = 'Issue Key'
                hdr_cells[2].text = 'Summary'
                hdr_cells[3].text = 'Status'
                
                # Add issues to table
                for issue in issues:
                    row_cells = issues_table.add_row().cells
                    row_cells[0].text = issue.get("issue type", "N/A")
                    row_cells[1].text = issue.get("issue key", "No key")
                    row_cells[2].text = issue.get("summary", "No summary")
                    row_cells[3].text = issue.get("status", "No status")

                # Add project summary section
                doc.add_heading("Project Summary", level=3)
                doc.add_paragraph(project_summary)
                doc.add_page_break()
                
            # Save document with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename_with_timestamp = f"{filename.split('.')[0]}_{timestamp}.docx"
            doc.save(filename_with_timestamp)
            print(f"Document saved successfully as {filename_with_timestamp}")
            
        except Exception as e:
            print(f"Error generating Word document: {e}")
    

    def extract_deliverables(self, projects_data: Dict[str, Dict]) -> Dict[str, List[Dict[str, str]]]:
        """
        Extract deliverable data using AI analysis, organized by project
        
        Uses Ollama LLM to intelligently identify key deliverables from JIRA issues
        for each project separately and returns them in project-organized format.
        
        Args:
            projects_data: Dictionary with project names as keys and project data as values
            
        Returns:
            Dictionary with project names as keys and lists of deliverable dictionaries as values
        """
        print("Analyzing projects to extract deliverables using AI...")
        
        project_deliverables = {}
        
        # Process each project separately for better organization
        for project_name, project_data in projects_data.items():
            print(f"Extracting deliverables for project: {project_name}")
            
            issues = project_data.get("issues", [])
            if not issues:
                project_deliverables[project_name] = []
                continue
            
            # Prepare project-specific data for AI analysis
            project_issues = []
            for issue in issues:
                project_issues.append({
                    "issue_type": issue.get("issue type", "Unknown"),
                    "issue_key": issue.get("issue key", "No key"),
                    "summary": issue.get("summary", "No summary"),
                    "status": issue.get("status", "Unknown"),
                    "due_date": issue.get("duedate", "No due date"),
                    "updated": issue.get("updated", "Unknown")
                })
            
            # Create project-specific AI analysis prompt
            prompt = self._create_project_deliverable_prompt(project_name, project_issues)
            
            try:
                # Call Ollama for AI analysis
                response = self._call_ollama_for_deliverables(prompt)
                
                if response:
                    deliverables = self._parse_deliverable_response(response, project_name)
                    if deliverables:
                        print(f"AI identified {len(deliverables)} deliverables for {project_name}")
                        project_deliverables[project_name] = deliverables
                        continue
                
                print(f"AI analysis failed for {project_name}, using fallback method")
                project_deliverables[project_name] = self._fallback_extract_project_deliverables(project_name, project_data)
                    
            except Exception as e:
                print(f"Error during AI deliverable extraction for {project_name}: {e}")
                project_deliverables[project_name] = self._fallback_extract_project_deliverables(project_name, project_data)
        
        return project_deliverables
    
    def _create_project_deliverable_prompt(self, project_name: str, project_issues: List[Dict]) -> str:
        """Create the AI prompt for project-specific deliverable extraction"""
        return f"""
You are a project management assistant. Analyze the following JIRA issues from project "{project_name}" and identify which ones represent key deliverables for this specific project.

A deliverable is typically:
- A work item that represents has been done and can show the outcome of this task.
- task's status is closed.
- Includes completed Stories, epics, tasks, or features (not bugs or sub-tasks).

Here are the issues from project "{project_name}" to analyze:
{json.dumps(project_issues, indent=2)}

Please return ONLY a valid JSON array of deliverables in this exact format:
[
  {{
    "deliverable_name": "Brief descriptive thing to deliverable",
    "due_date": "YYYY-MM-DD or original date format",
    "date_updated": "YYYY-MM-DD or original date format", 
    "status": "Status from JIRA"
  }}
]

"""
    
    def _call_ollama_for_deliverables(self, prompt: str) -> str:
        """Call Ollama API for deliverable analysis"""
        body = {
            "model": "llama3",
            "prompt": prompt,
            "stream": False
        }
        headers = {"Content-Type": "application/json"}
        
        print("Requesting AI analysis of deliverables...")
        response = requests.post(OLLAMA_URL, json=body, headers=headers)
        
        if response.status_code != 200:
            print(f"Ollama API error: {response.status_code} - {response.text}")
            return None
        
        result = response.json()
        return result.get("response", "").strip()
    
    def _parse_deliverable_response(self, ai_response: str, project_name: str) -> List[Dict[str, str]]:
        """Parse AI response and extract deliverable JSON for a specific project"""
        try:
            # Clean up the response to extract JSON
            json_start = ai_response.find('[')
            json_end = ai_response.rfind(']') + 1
            
            if json_start != -1 and json_end != -1:
                json_str = ai_response[json_start:json_end]
                deliverables = json.loads(json_str)
                
                # Format dates and add project name to deliverables
                for deliverable in deliverables:
                    deliverable["subproject"] = project_name
                    deliverable["due_date"] = self._format_date(deliverable.get("due_date", "No due date"))
                    deliverable["date_updated"] = self._format_date(deliverable.get("date_updated", "Unknown"))
                
                return deliverables
            else:
                print(f"Could not find valid JSON in AI response for {project_name}")
                return []
                
        except json.JSONDecodeError as e:
            print(f"Error parsing AI JSON response for {project_name}: {e}")
            print(f"AI Response: {ai_response[:500]}...")
            return []
    
    @staticmethod
    def _format_date(date_str: str) -> str:
        """Format date string to a more readable format"""
        if not date_str or date_str == "Unknown" or date_str == "No due date":
            return date_str
        try:
            # Parse ISO format and return just the date part
            if 'T' in date_str:
                return date_str.split('T')[0]
            return date_str
        except:
            return date_str
    
    def _fallback_extract_project_deliverables(self, project_name: str, project_data: Dict) -> List[Dict[str, str]]:
        """
        Fallback method to extract deliverables for a specific project using rule-based approach
        
        Args:
            project_name: Name of the project
            project_data: Dictionary containing project issues and data
            
        Returns:
            List of deliverable dictionaries for the specific project
        """
        print(f"Using fallback rule-based deliverable extraction for {project_name}...")
        
        deliverables = []
        deliverable_keywords = ["story", "epic", "task", "deliverable", "feature"]
        
        issues = project_data.get("issues", [])
        
        for issue in issues:
            issue_type = issue.get("issue type", "").lower()
            status = issue.get("status", "")
            
            # Filter for deliverable-type issues
            if any(keyword in issue_type for keyword in deliverable_keywords):
                deliverable = {
                    "subproject": project_name,
                    "deliverable_name": issue.get("summary", "No summary"),
                    "due_date": self._format_date(issue.get("duedate", "No due date")),
                    "date_updated": self._format_date(issue.get("updated", "Unknown")),
                    "status": status
                }
                deliverables.append(deliverable)
        
        return deliverables

    def generate_deliverable_table(self, doc: Document, project_deliverables: Dict[str, List[Dict[str, str]]]):
        """
        Generate deliverable overview table in the Word document, organized by project
        
        Args:
            doc: Document object to add the table to
            project_deliverables: Dictionary with project names as keys and lists of deliverables as values
        """
        # Count total deliverables across all projects
        total_deliverables = sum(len(deliverables) for deliverables in project_deliverables.values())
        
        if total_deliverables == 0:
            return
        
        # Add deliverable section heading
        doc.add_heading("Deliverables Overview", 2)
        
        print(f"Generating deliverable table with {total_deliverables} deliverables...")
        # Process each project's deliverables
        for project_name, deliverables in project_deliverables.items():
            if not deliverables:
                continue
                
            # Add project subheading
            doc.add_heading(f"{project_name} Deliverables", 3)
            
            # Create deliverable table for this project
            deliverable_table = doc.add_table(rows=1, cols=5)
            deliverable_table.style = 'Table Grid'
            
            # Set table headers (no need for subproject column since it's organized by project)
            hdr_cells = deliverable_table.rows[0].cells
            headers = ["Subproject", "Deliverable Name", "Due Date", "Date Updated", "Status"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # Add deliverables to table
            if deliverables and isinstance(deliverables, list):
                for deliverable in deliverables:
                    print(f"Adding deliverable to table: {deliverable}")
                    row_cells = deliverable_table.add_row().cells
                    row_cells[0].text = project_name
                    row_cells[1].text = str(deliverable.get("deliverable_name", "No name"))
                    row_cells[2].text = str(deliverable.get("due_date", "No due date"))
                    row_cells[3].text = str(deliverable.get("date_updated", "Unknown"))
                    row_cells[4].text = str(deliverable.get("status", "No status"))
            else:
                print(f"No valid deliverables found for project: {project_name}")
            
            doc.add_paragraph()  # Add spacing after each project's table

    def run(self):
        """
        Main execution method for processing multiple projects
        
        Orchestrates the complete workflow:
        1. Fetches issues from JIRA for each project
        2. Processes and formats issue data
        3. Generates AI summaries for each project
        4. Extracts deliverables using AI analysis
        5. Creates comprehensive Word document report
        """
        print("Starting JIRA to DOCX automation for multiple projects...")
        print("=" * 60)
        
        projects_data = {}
        
        # Process each project
        for project_name in self.project_names:
            print(f"\nProcessing project: {project_name}")
            print("-" * 40)
            
            # Step 1: Fetch issues from JIRA
            issues = self.fetch_issues(project_name)
            if not issues:
                print(f"No issues found for project {project_name}")
                projects_data[project_name] = {
                    "issues": [],
                    "ai_summary": f"No issues found for project {project_name}"
                }
                continue
            
            # Step 2: Process and format issue data
            issue_summaries = self._process_issues(issues)
            
            # Step 3: Generate AI summary for this project
            issues_string = self._format_issues_for_summary(issue_summaries)
            print(f"\nGenerating AI summary for {project_name}...")
            ai_summary = self.summarize_with_ollama(issues_string)
            print(f"AI Summary for {project_name}:\n{ai_summary}")
            
            # Store project data
            projects_data[project_name] = {
                "issues": issue_summaries,
                "ai_summary": ai_summary
            }
        
        # Step 4: Generate comprehensive Word document
        self._generate_final_report(projects_data)
        
    def _process_issues(self, issues: List[Dict]) -> List[Dict[str, str]]:
        """Process raw JIRA issues into formatted data"""
        issue_summaries = []
        
        for issue in issues:
            fields = issue.get("fields", {})
            issue_data = {
                "issue type": fields.get("issuetype", {}).get("name", "Unknown"),
                "issue key": issue.get("key", "No key"),
                "summary": fields.get("summary", "No summary"),
                "status": fields.get("status", {}).get("name", "Unknown"),
                "created": fields.get("created", "Unknown"),
                "updated": fields.get("updated", "Unknown"),
                "duedate": fields.get("duedate", "No due date"),
                "priority": fields.get("priority", {}).get("name", "None"),
            }
            issue_summaries.append(issue_data)
        
        return issue_summaries
    
    def _format_issues_for_summary(self, issue_summaries: List[Dict]) -> str:
        """Format issue data for AI summary generation"""
        return "\n".join([
            f"Issue Key: {issue['issue key']}, Summary: {issue['summary']}, "
            f"Status: {issue['status']}, Created: {issue['created']}, "
            f"Updated: {issue['updated']}, Due Date: {issue['duedate']}, "
            f"Priority: {issue['priority']}"
            for issue in issue_summaries
        ])
    
    def _generate_final_report(self, projects_data: Dict[str, Dict]):
        """Generate the final Word document report"""
        print(f"\n{'='*60}")
        print("Generating combined Word document for all projects...")
        
        total_issues = sum(len(project_data["issues"]) for project_data in projects_data.values())
        project_deliverables = self.extract_deliverables(projects_data)
        total_deliverables = sum(len(deliverables) for deliverables in project_deliverables.values())
        
        self.generate_word_document(projects_data)

        print(f"\n{'='*60}")
        print("Automation completed successfully!")
        print(f"Generated report for {len(projects_data)} projects with {total_issues} total issues.")
        print(f"Identified {total_deliverables} deliverables across all projects.")
        
        # Show deliverables breakdown by project
        for project_name, deliverables in project_deliverables.items():
            if deliverables:
                print(f"  - {project_name}: {len(deliverables)} deliverables")
    
 
def main():
    """
    Entry point for the JIRA automation application
    
    Initializes the automation system and handles configuration validation
    and error reporting for common setup issues.
    """
    try:
        project_names = JiraToDocxAutomation.from_env_projects()
        print(f"Project names loaded: {project_names}")
        
        automation = JiraToDocxAutomation(project_names)
        automation.run()
        
    except Exception as e:
        print(f"Error: {e}")
        print("\nPlease ensure:")
        print("1. Your .env file is properly configured")
        print("2. Ollama is running locally with llama3 model")
        print("3. Your JIRA credentials are correct")
        print("4. Project names exist in your JIRA instance")
        print("5. JIRA_PROJECTS is set in .env file (comma-separated)")


if __name__ == "__main__":
    main()
